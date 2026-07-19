import io

import pandas as pd
import streamlit as st

from utils.file_ops import init_module_state

PREFIX = "module5_word"


def run():
    st.title("📝 Word 处理及生成")
    st.caption("创建和编辑 Word 文档")

    init_module_state(PREFIX, {"content": ""})

    st.subheader("创建 Word 文档")
    
    doc_title = st.text_input("文档标题", value="报告文档", key=f"{PREFIX}_title")
    
    content = st.text_area(
        "文档内容",
        value=st.session_state.get(f"{PREFIX}_content", ""),
        key=f"{PREFIX}_content",
        height=300,
        placeholder="输入文档内容..."
    )

    col1, col2 = st.columns(2)
    with col1:
        font_size = st.slider("字体大小", min_value=10, max_value=24, value=12, key=f"{PREFIX}_font_size")
    with col2:
        font_name = st.selectbox("字体", ["宋体", "黑体", "楷体", "微软雅黑"], key=f"{PREFIX}_font_name")

    if st.button("生成 Word 文档", key=f"{PREFIX}_generate", type="primary"):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn

            doc = Document()
            
            style = doc.styles['Normal']
            style.font.name = font_name
            style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            style.font.size = Pt(font_size)

            title_heading = doc.add_heading(doc_title, level=0)
            for run in title_heading.runs:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            st.success("Word 文档生成成功！")
            st.download_button(
                label="下载 Word 文档",
                data=output.getvalue(),
                file_name=f"{doc_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"{PREFIX}_download",
            )
        except ImportError:
            st.warning("需要安装 python-docx 库，请运行 pip install python-docx")
        except Exception as e:
            st.error(f"生成失败：{str(e)}")

    st.subheader("Word 模板")
    st.info("💡 提示：后续可扩展支持上传 Word 模板并填充数据")